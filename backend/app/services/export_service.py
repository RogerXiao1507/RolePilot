from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.schemas.full_resume import FullTailoredResumeDraftResponse


EXPORT_DIR = Path("generated_exports")
EXPORT_DIR.mkdir(exist_ok=True)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_section_heading(doc: Document, title: str):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=3, line=1.0)

    run = p.add_run(title.upper())
    set_run_font(run, size=12, bold=True)

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_pr = p._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "444444")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style=None)
    p.style = doc.styles["List Bullet"]
    set_paragraph_spacing(p, before=0, after=0, line=1.3)

    run = p.add_run(text)
    set_run_font(run, size=10)


def add_entry_header(
    doc: Document,
    left_top: str,
    left_bottom: str | None,
    right_top: str | None,
    right_bottom: str | None,
):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(5.15)
    table.columns[1].width = Inches(2.35)

    for row in table.rows:
        row.cells[0].width = Inches(5.15)
        row.cells[1].width = Inches(2.35)

    # Left cell
    left_cell = table.cell(0, 0)
    left_cell.text = ""

    left_p1 = left_cell.paragraphs[0]
    left_p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(left_p1, before=0, after=0, line=1.0)
    run = left_p1.add_run(left_top)
    set_run_font(run, size=12, bold=True)

    if left_bottom:
        left_p2 = left_cell.add_paragraph()
        left_p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(left_p2, before=0, after=0, line=1.0)
        run = left_p2.add_run(left_bottom)
        set_run_font(run, size=10, italic=True)

    # Right cell
    right_cell = table.cell(0, 1)
    right_cell.text = ""

    right_p1 = right_cell.paragraphs[0]
    right_p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(right_p1, before=0, after=0, line=1.0)
    if right_top:
        run = right_p1.add_run(right_top)
        set_run_font(run, size=12, bold=True)

    if right_bottom:
        right_p2 = right_cell.add_paragraph()
        right_p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(right_p2, before=0, after=0, line=1.0)
        run = right_p2.add_run(right_bottom)
        set_run_font(run, size=10, italic=True)

    # Remove table borders
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ["top", "left", "bottom", "right"]:
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "nil")
                tc_borders.append(el)
            tc_pr.append(tc_borders)

def add_project_header(
    doc: Document,
    title: str,
    date_range: str | None,
    location: str | None,
):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(5.15)
    table.columns[1].width = Inches(2.35)

    for row in table.rows:
        row.cells[0].width = Inches(5.15)
        row.cells[1].width = Inches(2.35)

    # Left cell
    left_cell = table.cell(0, 0)
    left_cell.text = ""

    left_p1 = left_cell.paragraphs[0]
    left_p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(left_p1, before=0, after=0, line=1.0)
    run = left_p1.add_run(title)
    set_run_font(run, size=12, bold=True)

    # Right cell
    right_cell = table.cell(0, 1)
    right_cell.text = ""

    right_p1 = right_cell.paragraphs[0]
    right_p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(right_p1, before=0, after=0, line=1.0)
    if date_range:
        run = right_p1.add_run(date_range)
        set_run_font(run, size=12, bold=True)

    if location:
        right_p2 = right_cell.add_paragraph()
        right_p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(right_p2, before=0, after=0, line=1.0)
        run = right_p2.add_run(location)
        set_run_font(run, size=10, italic=True)

    # Remove table borders
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ["top", "left", "bottom", "right"]:
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "nil")
                tc_borders.append(el)
            tc_pr.append(tc_borders)

def build_tailored_resume_docx(
    *,
    draft: FullTailoredResumeDraftResponse,
    output_path: str,
):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(10)

    # Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name_p, before=0, after=2, line=1.0)
    run = name_p.add_run(draft.header.name)
    set_run_font(run, size=26, bold=True)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact_p, before=0, after=8, line=1.2)
    contact_text = " | ".join(
        [x for x in [draft.header.location, draft.header.phone, draft.header.email, *draft.header.websites] if x]
    )
    run = contact_p.add_run(contact_text)
    set_run_font(run, size=10)

    # Education
    add_section_heading(doc, "Education")
    for edu in draft.education:
        add_entry_header(
            doc,
            edu.school,
            edu.degree,
            edu.date_range,
            edu.location,
        )
        if edu.gpa:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=0, after=0, line=1.2)
            r1 = p.add_run("GPA: ")
            set_run_font(r1, size=10, bold=False)
            r2 = p.add_run(edu.gpa)
            set_run_font(r2, size=10)

        if edu.coursework:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=0, after=4, line=1.2)
            r1 = p.add_run("Coursework: ")
            set_run_font(r1, size=10, bold=True)
            r2 = p.add_run(", ".join(edu.coursework))
            set_run_font(r2, size=10)

    # Experience
    add_section_heading(doc, "Experience")
    for exp in draft.experience:
        add_entry_header(
            doc,
            exp.title,
            exp.subtitle,
            exp.date_range,
            exp.location,
        )
        for bullet in exp.bullets:
            add_bullet(doc, bullet)

    # Projects
    add_section_heading(doc, "Projects")
    for proj in draft.projects:
        add_project_header(
            doc,
            proj.title,
            proj.date_range,
            proj.location,
        )
        for bullet in proj.bullets:
            add_bullet(doc, bullet)

    # Skills
    add_section_heading(doc, "Skills")
    skill_lines = [
        ("Programming Languages", draft.skills.programming_languages),
        ("Frameworks / Tools", draft.skills.frameworks_tools),
        ("Hardware / Instrumentation", draft.skills.hardware_instrumentation),
        ("Technical Areas", draft.skills.technical_areas),
        ("Developer Tools", draft.skills.developer_tools),
    ]
    for label, values in skill_lines:
        if not values:
            continue
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line=1.2)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=10, bold=True)
        r2 = p.add_run(", ".join(values))
        set_run_font(r2, size=10)

    doc.save(output_path)
    return output_path